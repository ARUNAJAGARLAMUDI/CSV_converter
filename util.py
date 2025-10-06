#utils.py
 
 
import pandas as pd
from docx import Document
from io import BytesIO
 
 
# Rule-based summary generator
def generate_summary(row: pd.Series) -> str:
    desc = (row.get('description') or '').strip()
    short_desc = (row.get('short_description') or '').strip()
    affected = (row.get('affected_customers') or 'N/A')
    state = (row.get('state') or 'N/A')
    code = (row.get('completion_code') or 'N/A')
 
    lead = short_desc or (desc[:120] + '...' if desc else 'the work described')
 
    summary = f"What is this project? {lead}. "
    summary += f"Who does it help? It helps {affected}. "
    summary += f"What's happening now? The project is '{state}' and has a code '{code}'. "
 
    # Extra simple explanations for non-technical readers
    lower = desc.lower()
    if any(k in lower for k in ['upgrade', 'migrate', 'migration']):
        summary += "Put simply: We are making things better or moving them so they work faster and easier for everyone."
    elif any(k in lower for k in ['fix', 'bug', 'issue', 'resolve']):
        summary += "Put simply: We are fixing something that's broken so people have fewer problems."
    elif any(k in lower for k in ['implement', 'deploy', 'build']):
        summary += "Put simply: We are adding something new that people can use in their daily work."
    elif any(k in lower for k in ['compliance', 'validation', 'calibration']):
        summary += "Put simply: We are making sure everything is safe and follows the rules."
    else:
        summary += "Put simply: This project is here to make things easier or better for the people it helps."
 
    return summary
 
 
 
def create_docx(row: pd.Series) -> BytesIO:
    doc = Document()
    p_number = row.get('p_number', 'UNKNOWN')
    short_desc = row.get('short_description', '')
    desc = row.get('description', '')
    affected = row.get('affected_customers', 'N/A')
 
    doc.add_heading(f"Project Summary – {p_number}", level=0)
    if short_desc:
        doc.add_heading(str(short_desc), level=1)
    doc.add_paragraph(f"Project Number: {p_number}")
    if desc:
        doc.add_paragraph(f"Description: {desc}")
    doc.add_paragraph(generate_summary(row))
    doc.add_paragraph(f"Affected Customers: {affected}")
 
    # Add technical terms explanation section
    doc.add_heading("Technical Terms Explained", level=1)
    technical_terms = {
        "migration": "Migration means moving data or systems from one place to another. For example, moving files from an old computer to a new one.",
        "upgrade": "Upgrade means improving something to a newer or better version. For example, updating your phone's software.",
        "compliance": "Compliance means following rules or standards set by authorities. For example, making sure a product meets safety regulations.",
        "deployment": "Deployment means putting a new system or feature into use. For example, launching a new website.",
        "validation": "Validation means checking if something works as expected. For example, testing if a new app runs correctly.",
        "calibration": "Calibration means adjusting something to make it accurate. For example, setting a scale to show the correct weight.",
        "bug": "A bug is an error or flaw in a system that causes it to behave unexpectedly. For example, an app crashing when you press a button.",
        "issue": "An issue is a problem that needs to be fixed. For example, slow internet speed at home.",
        "implementation": "Implementation means putting a plan or idea into action. For example, starting a recycling program at school."
    }
    for term, explanation in technical_terms.items():
        doc.add_paragraph(f"{term.capitalize()}: {explanation}")
 
    # Add detailed layman summary (about 150 words)
    doc.add_heading("Project Details for Everyone", level=1)
    details = f"This project addresses several important challenges. The main problem is described as: {desc if desc else 'No description provided.'} The team is working to resolve these issues by following a structured approach, which includes planning, testing, and implementing solutions. Progress so far includes identifying the root causes, starting necessary upgrades or migrations, and ensuring compliance with all required standards. Technical terms like migration, upgrade, and compliance are explained above to help everyone understand the process. The project is currently in the '{row.get('state', 'N/A')}' stage, and the completion code is '{row.get('completion_code', 'N/A')}'. The goal is to make things easier and better for all affected customers. If you have questions about any part of the project, please refer to the explanations above or contact the project team for more details."
    doc.add_paragraph(details)
 
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)  # Important: reset buffer position
    return buf
 
 
def create_combined_docx(df: pd.DataFrame) -> BytesIO:
    doc = Document()
    technical_terms = {
        "migration": "Migration means moving data or systems from one place to another. For example, moving files from an old computer to a new one.",
        "upgrade": "Upgrade means improving something to a newer or better version. For example, updating your phone's software.",
        "compliance": "Compliance means following rules or standards set by authorities. For example, making sure a product meets safety regulations.",
        "deployment": "Deployment means putting a new system or feature into use. For example, launching a new website.",
        "validation": "Validation means checking if something works as expected. For example, testing if a new app runs correctly.",
        "calibration": "Calibration means adjusting something to make it accurate. For example, setting a scale to show the correct weight.",
        "bug": "A bug is an error or flaw in a system that causes it to behave unexpectedly. For example, an app crashing when you press a button.",
        "issue": "An issue is a problem that needs to be fixed. For example, slow internet speed at home.",
        "implementation": "Implementation means putting a plan or idea into action. For example, starting a recycling program at school."
    }
    for idx, row in df.iterrows():
        p_number = row.get('p_number', 'UNKNOWN')
        short_desc = row.get('short_description', '')
        desc = row.get('description', '')
        affected = row.get('affected_customers', 'N/A')
        doc.add_heading(f"Project Summary – {p_number}", level=0)
        if short_desc:
            doc.add_heading(str(short_desc), level=1)
        doc.add_paragraph(f"Project Number: {p_number}")
        if desc:
            doc.add_paragraph(f"Description: {desc}")
        doc.add_paragraph(generate_summary(row))
        doc.add_paragraph(f"Affected Customers: {affected}")
 
        # Add technical terms explanation section
        doc.add_heading("Technical Terms Explained", level=1)
        for term, explanation in technical_terms.items():
            doc.add_paragraph(f"{term.capitalize()}: {explanation}")
 
        # Add detailed layman summary (about 150 words)
        doc.add_heading("Project Details for Everyone", level=1)
        details = f"This project addresses several important challenges. The main problem is described as: {desc if desc else 'No description provided.'} The team is working to resolve these issues by following a structured approach, which includes planning, testing, and implementing solutions. Progress so far includes identifying the root causes, starting necessary upgrades or migrations, and ensuring compliance with all required standards. Technical terms like migration, upgrade, and compliance are explained above to help everyone understand the process. The project is currently in the '{row.get('state', 'N/A')}' stage, and the completion code is '{row.get('completion_code', 'N/A')}'. The goal is to make things easier and better for all affected customers. If you have questions about any part of the project, please refer to the explanations above or contact the project team for more details."
        doc.add_paragraph(details)
 
        if idx != df.index[-1]:
            doc.add_page_break()
 
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)  # Important: reset buffer position
    return buf